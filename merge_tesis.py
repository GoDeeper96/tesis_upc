# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

CAP1 = 'D:/Proyectos actuales/proyecto_tesis_upc/Capitulo1_Tesis_Final.docx'
CAP2 = 'D:/Proyectos actuales/proyecto_tesis_upc/Capitulo2_Marcoteorico_v3.docx'
OUT  = 'D:/Proyectos actuales/proyecto_tesis_upc/Tesis_Final.docx'


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def append_doc(base, source_path):
    src = Document(source_path)

    # Copy styles that exist in source but not in base
    src_style_ids = {s.style_id for s in src.styles}
    base_style_ids = {s.style_id for s in base.styles}
    for style in src.styles:
        if style.style_id not in base_style_ids and style.element is not None:
            try:
                base.styles.element.append(deepcopy(style.element))
            except Exception:
                pass

    # Copy all body children except final sectPr
    body = src.element.body
    children = list(body)
    for elem in children:
        if elem.tag == qn('w:sectPr'):
            continue
        base.element.body.append(deepcopy(elem))


base_doc = Document(CAP1)
add_page_break(base_doc)
append_doc(base_doc, CAP2)
base_doc.save(OUT)
print(f'Saved: {OUT}')
