"""
Generate Project_Charter_Tesis.docx
Following the exact structure of Project_Charter_ejemplo.docx (MINTRA)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
s = doc.sections[0]
s.top_margin    = Cm(2.54)
s.bottom_margin = Cm(2.54)
s.left_margin   = Cm(3.0)
s.right_margin  = Cm(2.54)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def header_row(table, cols, color='1F3864'):
    row = table.rows[0]
    for i, col in enumerate(cols):
        c = row.cells[i]
        c.text = col
        r = c.paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(c, color)

def centered(text, bold=False, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p

def bold_inline(label, text):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
centered('UNIVERSIDAD PERUANA DE CIENCIAS APLICADAS', bold=True, size=14)
doc.add_paragraph()
centered(
    'Solución tecnológica basada en IA para optimizar el tiempo en la\n'
    'generación de contenido de cursos virtuales en universidades del Perú',
    bold=False, size=13
)
doc.add_paragraph()
centered('Project Charter', bold=True, size=13)
centered('Versión 1.0', bold=False, size=11)
doc.add_paragraph()
doc.add_paragraph()
body('Preparado por:', bold=False)
body('[Nombre Apellido] — [Rol]')
body('[Nombre Apellido] — [Rol]')
doc.add_paragraph()
body('Lima, abril de 2026')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# HISTORIAL DE REVISIONES
# ═══════════════════════════════════════════════════════════════════════════════
body('Historial de Revisiones', bold=True)
rev_tbl = doc.add_table(rows=2, cols=4)
rev_tbl.style = 'Table Grid'
header_row(rev_tbl, ['Versión', 'Fecha', 'Autor', 'Descripción'])
rev_tbl.rows[1].cells[0].text = '1'
rev_tbl.rows[1].cells[1].text = 'Abril 2026'
rev_tbl.rows[1].cells[2].text = '[Nombre Apellido]'
rev_tbl.rows[1].cells[3].text = 'Versión inicial del Project Charter'
doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# RESUMEN EJECUTIVO
# ═══════════════════════════════════════════════════════════════════════════════
h1('Resumen ejecutivo')

body(
    'Las universidades peruanas con oferta académica virtual enfrentan una demora '
    'estructural en la producción de material educativo para sus cursos. Esta demora '
    'se concentra en la fase de generación del contenido instruccional denominado IPES '
    '(Introducción, Presentaciones y Ejercicios), que actualmente es elaborado '
    'manualmente por el docente virtualizador —un profesional especializado que redacta '
    'guiones de narración, introducciones semanales y consignas de actividades—, con '
    'tiempos de producción que oscilan entre dos y tres semanas por curso y costos '
    'operativos que escalan proporcionalmente con el volumen de cursos.'
)
body(
    'Este documento detalla el proyecto de titulación que propone la implementación '
    'de una solución tecnológica basada en agentes de inteligencia artificial para '
    'optimizar el tiempo y la generación de contenido de los cursos virtuales. La '
    'solución automatiza el proceso en tres fases: diseño del esquema del curso, '
    'aseguramiento pedagógico y generación del contenido textual, entregando como '
    'resultado un documento Word por semana para las 18 semanas que comprende un '
    'curso completo. La efectividad del contenido generado es evaluada por docentes '
    'revisores mediante métricas de calidad pedagógica.'
)
body(
    'El proyecto se alinea con el objetivo general aprobado por el asesor académico: '
    'implementar la solución tecnológica basada en inteligencia artificial para '
    'optimizar el tiempo, la generación de contenido y los cursos virtuales en la '
    'universidad, ejecutado en cuatro fases que cubren el análisis, diseño, '
    'implementación/validación y plan de continuidad de la solución.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# MARCO TEÓRICO
# ═══════════════════════════════════════════════════════════════════════════════
h1('Marco Teórico')

terms = [
    ('IPES (Introducción, Presentaciones y Ejercicios): ',
     'Conjunto de materiales instruccionales textuales que componen el contenido de '
     'cada semana académica de un curso virtual. Incluye la introducción a la semana, '
     'los guiones de narración para los recursos de presentación y las consignas de '
     'los ejercicios, tanto calificados como no calificados.'),
    ('Docente Virtualizador: ',
     'Profesional especializado responsable de redactar manualmente el contenido '
     'textual de los materiales IPES, asegurando su alineación con los logros de '
     'aprendizaje del sílabo y los criterios pedagógicos institucionales.'),
    ('Docente Revisor: ',
     'Profesional que valida la calidad, coherencia pedagógica y alineación del '
     'contenido instruccional generado, ya sea producido manualmente por el docente '
     'virtualizador o generado por los agentes de IA.'),
    ('Modelo de Lenguaje de Gran Escala (LLM): ',
     'Modelo de inteligencia artificial entrenado sobre grandes volúmenes de texto '
     'que tiene la capacidad de generar, resumir, traducir y razonar sobre lenguaje '
     'natural. En el presente proyecto se utilizan GPT-4.1-mini (OpenAI) y '
     'Gemini 2.0-flash-lite (Google) como motores de generación de contenido.'),
    ('Agente de IA: ',
     'Sistema autónomo basado en LLMs que percibe su entorno, toma decisiones y '
     'ejecuta acciones para lograr un objetivo específico. En este proyecto, los '
     'agentes operan en cadena para diseñar, validar y generar el contenido de '
     'cada semana del curso.'),
    ('LangGraph: ',
     'Framework de orquestación de agentes de IA basado en grafos de estado que '
     'permite definir flujos de trabajo complejos con nodos de decisión, reintentos '
     'y validaciones. Es la base tecnológica sobre la que se implementan los agentes '
     'del presente proyecto.'),
    ('Diseño Instruccional: ',
     'Disciplina que se ocupa del análisis, diseño, desarrollo, implementación y '
     'evaluación de experiencias de aprendizaje. El modelo ADDIE (Analysis, Design, '
     'Development, Implementation, Evaluation) es el marco de referencia institucional '
     'sobre el que se estructura el proceso de virtualización de cursos.'),
    ('Taxonomía de Bloom: ',
     'Marco de clasificación de objetivos de aprendizaje organizado en niveles '
     'cognitivos (Recordar, Comprender, Aplicar, Analizar, Evaluar, Crear). En el '
     'proceso de diseño instruccional, el verbo del logro de aprendizaje define el '
     'tipo de recurso a utilizar en cada semana.'),
    ('Scrum: ',
     'Metodología ágil de gestión de proyectos basada en iteraciones cortas '
     'denominadas sprints, que permite una entrega incremental y adaptativa de '
     'los productos del proyecto.'),
]

for label, definition in terms:
    bold_inline(label, definition)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# POSICIONAMIENTO
# ═══════════════════════════════════════════════════════════════════════════════
h1('Posicionamiento')

body(
    'La demora en la gestión de contenido y materiales de los cursos virtuales en la '
    'institución constituye el problema central del proyecto. La producción de material '
    'educativo para los cursos virtuales depende de la disponibilidad del docente '
    'virtualizador, cuya labor manual no puede escalar al ritmo de la creciente demanda '
    'de cursos virtuales, generando retrasos y costos operativos elevados. A continuación '
    'se sintetiza la problemática:'
)

pos_tbl = doc.add_table(rows=2, cols=2)
pos_tbl.style = 'Table Grid'
header_row(pos_tbl, ['Problema', 'Causas'])
pos_tbl.rows[1].cells[0].text = (
    'Demora en la producción de material educativo para los cursos virtuales, derivada '
    'de la dependencia del proceso en el docente virtualizador y la ausencia de '
    'herramientas tecnológicas que optimicen la generación de contenido instruccional.'
)
pos_tbl.rows[1].cells[1].text = (
    '• Dependencia exclusiva del docente virtualizador para producir el contenido '
    'semanal de cada curso (~2 a 3 semanas por curso).\n'
    '• Remuneración por hora que escala linealmente con el volumen de cursos, '
    'incrementando los costos operativos sin mecanismo de reducción.\n'
    '• Ausencia de herramientas de IA en la fase de generación de contenido '
    'instruccional.\n'
    '• Proceso no escalable que representa un riesgo operativo a largo plazo para '
    'instituciones con alta demanda de cursos virtuales.'
)
doc.add_paragraph()

# ── Objetivos ─────────────────────────────────────────────────────────────────
body('Objetivos', bold=True)
doc.add_paragraph()
body('Objetivo General', bold=True)
body(
    'Implementar la solución tecnológica basada en inteligencia artificial para '
    'optimizar el tiempo, la generación de contenido y los cursos virtuales en '
    'la universidad.'
)
doc.add_paragraph()
body('Objetivos Específicos', bold=True)

oes = [
    ('OE1: ', 'Analizar la solución tecnológica basada en IA para optimizar el tiempo '
              'en la generación de contenido de cursos virtuales.'),
    ('OE2: ', 'Diseñar la arquitectura de la solución tecnológica basada en IA para '
              'optimizar el tiempo en la generación de contenido de cursos virtuales '
              'en universidades del Perú.'),
    ('OE3: ', 'Implementar y validar la solución tecnológica basada en IA para '
              'optimizar el tiempo en la generación de contenido de cursos virtuales '
              'mediante métricas de rendimiento.'),
    ('OE4: ', 'Elaborar un plan de continuidad que asegure la viabilidad técnica de '
              'la solución tecnológica propuesta en el tiempo.'),
]
for label, text in oes:
    bold_inline(label, text)

doc.add_paragraph()

# ── Indicadores de éxito ──────────────────────────────────────────────────────
body('Indicadores de Éxito', bold=True)
body(
    'Los indicadores de éxito permiten precisar el nivel de logro esperado en cada '
    'objetivo específico. La siguiente tabla relaciona cada indicador con su objetivo '
    'correspondiente.'
)

ie_tbl = doc.add_table(rows=5, cols=3)
ie_tbl.style = 'Table Grid'
header_row(ie_tbl, ['Indicador', 'Descripción', 'Objetivo'])
ie_data = [
    ('IE1', 'Aprobación del análisis de tecnologías y enfoques de IA por el asesor académico.', 'OE1'),
    ('IE2', 'Aprobación de la arquitectura de agentes de IA por el asesor académico.', 'OE2'),
    ('IE3', 'Evaluación favorable del contenido generado mediante métricas de rendimiento y revisión de docentes expertos.', 'OE3'),
    ('IE4', 'Aprobación del plan de continuidad por el asesor académico.', 'OE4'),
]
for i, (code, desc, obj) in enumerate(ie_data):
    row = ie_tbl.rows[i+1]
    row.cells[0].text = code
    row.cells[1].text = desc
    row.cells[2].text = obj

doc.add_paragraph()
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ALCANCE DEL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
h1('Alcance del proyecto')

body(
    'El proyecto aborda la fase de generación de contenido textual del proceso de '
    'virtualización de cursos, utilizando agentes de IA para producir los materiales '
    'IPES de las 18 semanas de un curso completo. El alcance contempla los siguientes puntos:'
)

body('El proyecto incluirá:', bold=True)
includes = [
    'Elaboración del Project Charter y documentación de inicio del proyecto.',
    'Análisis de tecnologías LLM y frameworks de agentes para la generación de contenido instruccional.',
    'Diseño de la arquitectura multi-agente con fases de diseño, aseguramiento y generación de contenido.',
    'Implementación de los agentes de IA: KickOffAgent, SyllabusActivityClassifier, CourseSchemaAgent, IpesAgent y CheckIpesAgent.',
    'Generación del contenido IPES (introducción, guiones de presentación y consignas de ejercicios) para cursos de prueba.',
    'Evaluación del contenido generado por docentes revisores mediante rúbrica de calidad pedagógica.',
    'Análisis comparativo de tiempo y costo de producción respecto al proceso manual.',
    'Elaboración del plan de continuidad con roles de soporte, niveles de servicio y gestión de riesgos.',
    'Documentación técnica: informe de análisis, arquitectura, resultados de validación y plan de continuidad.',
]
for item in includes:
    bullet(item)

doc.add_paragraph()
body('El proyecto excluirá:', bold=True)
excludes = [
    'Desarrollo o implementación de plataformas de gestión de cursos o LMS.',
    'Producción de recursos multimedia (video, audio, presentaciones, H5P, Storyline).',
    'Automatización de otras etapas del proceso de virtualización fuera de la generación de contenido textual.',
    'Integración con sistemas institucionales externos (ERP, LMS, repositorios).',
    'Mantenimiento de la solución posterior a la entrega del trabajo de investigación.',
    'Implementación en producción a escala institucional.',
]
for item in excludes:
    bullet(item)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# RESTRICCIONES
# ═══════════════════════════════════════════════════════════════════════════════
h1('Restricciones')

body(
    'Las restricciones siguientes buscan mantener el proyecto enfocado en el '
    'cumplimiento del objetivo general.'
)

rest_tbl = doc.add_table(rows=6, cols=2)
rest_tbl.style = 'Table Grid'
header_row(rest_tbl, ['Restricción', 'Descripción'])
rest_data = [
    ('Tiempo',
     'El proyecto debe concluirse dentro del ciclo académico 2026 del Programa de '
     'Titulación de la UPC.'),
    ('Alcance',
     'La solución cubre exclusivamente la generación de contenido textual IPES; '
     'no contempla producción multimedia ni integración con plataformas.'),
    ('Costo',
     'El presupuesto de IA por curso no debe superar los $2.00 USD por ejecución '
     'completa (18 semanas), para mantener viabilidad económica frente al proceso manual.'),
    ('Calidad',
     'Todo el contenido generado debe ser evaluado por al menos un docente revisor '
     'antes de considerarse válido para efectos del proyecto.'),
    ('Confidencialidad',
     'Los documentos institucionales utilizados como entrada (sílabos, actas, '
     'bibliografías) deben tratarse con reserva y no divulgarse fuera del contexto '
     'académico del proyecto.'),
]
for i, (name, desc) in enumerate(rest_data):
    row = rest_tbl.rows[i+1]
    row.cells[0].text = name
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = desc

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# IMPACTO EN LA ORGANIZACIÓN
# ═══════════════════════════════════════════════════════════════════════════════
h1('Impacto en la organización')

body(
    'La implementación de la solución tecnológica basada en agentes de IA para la '
    'generación de contenido de cursos virtuales busca generar los siguientes impactos '
    'en la institución de educación superior:'
)

impacts = [
    'Reducir el tiempo de producción de contenido IPES de dos a tres semanas a pocas '
    'horas de generación más revisión, lo que permite acelerar el lanzamiento de '
    'cursos virtuales al estudiante.',
    'Disminuir los costos operativos asociados a la remuneración por horas del '
    'docente virtualizador, al reducir la carga de trabajo manual en la fase de '
    'generación de contenido.',
    'Homogenizar la calidad del contenido instruccional entre cursos, al aplicar '
    'las mismas reglas pedagógicas a todos los cursos procesados por el sistema.',
    'Escalar la capacidad de producción de cursos virtuales sin incrementar '
    'proporcionalmente los recursos humanos asignados a la fase de generación.',
    'Sentar las bases para la integración futura de la generación de contenido '
    'con otras etapas del flujo de virtualización que podrían ser asistidas por IA.',
]
for item in impacts:
    bullet(item)

doc.add_paragraph()
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ORGANIZACIÓN DEL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
h1('Organización del Proyecto')

body('Equipo del Proyecto', bold=True)
body('El equipo del proyecto estará conformado por:')

team_tbl = doc.add_table(rows=4, cols=3)
team_tbl.style = 'Table Grid'
header_row(team_tbl, ['Rol', 'Miembro', 'Responsabilidades'])
team_data = [
    ('Project Manager / Investigador Principal',
     '[Nombre Apellido]',
     'Liderar el desarrollo del proyecto, coordinar las fases de análisis, diseño, '
     'implementación y validación. Gestionar los entregables y la comunicación con '
     'el asesor académico.'),
    ('Desarrollador IA / Co-investigador',
     '[Nombre Apellido]',
     'Implementar los agentes de IA (LangGraph, FastAPI), integrar los modelos LLM '
     '(GPT-4.1-mini, Gemini 2.0-flash-lite) y ejecutar las pruebas de generación.'),
    ('Asesor Académico',
     'Daniel Burga',
     'Orientar el desarrollo del trabajo de suficiencia, revisar y aprobar los '
     'entregables de cada fase del proyecto.'),
]
for i, (rol, member, resp) in enumerate(team_data):
    row = team_tbl.rows[i+1]
    row.cells[0].text = rol
    row.cells[1].text = member
    row.cells[2].text = resp

doc.add_paragraph()

body('Stakeholders', bold=True)
body(
    'Los stakeholders son personas u organizaciones con interés en el resultado del '
    'proyecto, ya sea de manera directa o indirecta.'
)

sk_tbl = doc.add_table(rows=5, cols=3)
sk_tbl.style = 'Table Grid'
header_row(sk_tbl, ['Stakeholder', 'Tipo', 'Interés / Necesidad'])
sk_data = [
    ('Institución de educación superior (universidad)',
     'Externo',
     'Reducir tiempos y costos en la producción de cursos virtuales; escalar '
     'la oferta académica virtual con eficiencia operativa.'),
    ('Docentes Revisores',
     'Externo',
     'Recibir contenido instruccional de calidad para su revisión, alineado '
     'a los criterios pedagógicos institucionales.'),
    ('Docentes Virtualizadores',
     'Externo',
     'Reducir la carga de trabajo manual en la redacción de materiales IPES, '
     'enfocándose en tareas de mayor valor pedagógico.'),
    ('Equipo de diseño instruccional (LXD)',
     'Externo',
     'Agilizar el proceso de producción de cursos y mejorar la consistencia '
     'del contenido generado.'),
]
for i, (name, tipo, interest) in enumerate(sk_data):
    row = sk_tbl.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = tipo
    row.cells[2].text = interest

doc.add_paragraph()
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# RECURSOS REQUERIDOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('Recursos requeridos')

body('Recursos Humanos', bold=True)
hum_tbl = doc.add_table(rows=4, cols=3)
hum_tbl.style = 'Table Grid'
header_row(hum_tbl, ['Rol', 'Horas estimadas', 'Responsabilidad principal'])
hum_data = [
    ('Investigador / Project Manager', '~300 horas', 'Gestión, documentación y coordinación del proyecto.'),
    ('Desarrollador IA', '~400 horas', 'Implementación de agentes, pruebas y validación técnica.'),
    ('Asesor Académico', '~30 horas', 'Revisión y aprobación de entregables.'),
]
for i, row_data in enumerate(hum_data):
    row = hum_tbl.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val

doc.add_paragraph()
body('Hardware', bold=True)
hw_tbl = doc.add_table(rows=3, cols=2)
hw_tbl.style = 'Table Grid'
header_row(hw_tbl, ['Recurso', 'Descripción'])
hw_data = [
    ('Laptop de desarrollo', 'Computadora de alto rendimiento para implementar y ejecutar los agentes de IA, con acceso a internet para consumir APIs de LLM.'),
    ('Acceso a internet', 'Conexión estable para el consumo de APIs de OpenAI (GPT-4.1-mini) y Google (Gemini 2.0-flash-lite).'),
]
for i, (name, desc) in enumerate(hw_data):
    row = hw_tbl.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = desc

doc.add_paragraph()
body('Software', bold=True)
sw_tbl = doc.add_table(rows=9, cols=2)
sw_tbl.style = 'Table Grid'
header_row(sw_tbl, ['Recurso', 'Descripción'])
sw_data = [
    ('Python 3.11+', 'Lenguaje de programación principal para el backend y los agentes de IA.'),
    ('FastAPI', 'Framework para exponer los agentes como servicios REST.'),
    ('LangGraph', 'Framework de orquestación de agentes de IA basado en grafos de estado.'),
    ('GPT-4.1-mini (OpenAI API)', 'Modelo LLM para lectura de documentos, generación de esquemas y clasificación.'),
    ('Gemini 2.0-flash-lite (Google API)', 'Modelo LLM de bajo costo para generación de contenido IPES.'),
    ('Next.js + React 18', 'Frontend para la interfaz de usuario de la solución.'),
    ('Visual Studio Code', 'IDE principal para el desarrollo del proyecto.'),
    ('Microsoft Word / python-docx', 'Herramienta para la generación y revisión de documentos IPES por semana.'),
]
for i, (name, desc) in enumerate(sw_data):
    row = sw_tbl.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = desc

doc.add_paragraph()
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# FASES E HITOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('Fases e hitos del proyecto')

body(
    'A continuación se presentan las fases del proyecto, sus fechas estimadas de '
    'culminación, los entregables incluidos y la prioridad de cada hito.'
)

hitos_tbl = doc.add_table(rows=5, cols=4)
hitos_tbl.style = 'Table Grid'
header_row(hitos_tbl, ['Hito', 'Fecha estimada', 'Entregables', 'Prioridad'])
hitos_data = [
    ('Aprobación del Project Charter y análisis (OE1)',
     'Mayo 2026',
     'Project Charter aprobado.\nInforme de análisis de tecnologías LLM y frameworks de agentes.',
     'Alta'),
    ('Aprobación del diseño de arquitectura (OE2)',
     'Junio 2026',
     'Arquitectura multi-agente documentada (fases: diseño, aseguramiento, generación).\nDiagramas de flujo del proceso.',
     'Alta'),
    ('Aprobación de implementación y validación (OE3)',
     'Agosto 2026',
     'Agentes de IA implementados y funcionales.\nResultados de evaluación por docentes revisores.\nAnálisis comparativo tiempo/costo vs. proceso manual.',
     'Alta'),
    ('Aprobación del plan de continuidad (OE4)',
     'Septiembre 2026',
     'Plan de continuidad con roles de soporte, niveles de servicio, gestión de riesgos y costos.',
     'Media'),
]
for i, (hito, fecha, entregables, prioridad) in enumerate(hitos_data):
    row = hitos_tbl.rows[i+1]
    row.cells[0].text = hito
    row.cells[1].text = fecha
    row.cells[2].text = entregables
    row.cells[3].text = prioridad

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# ENFOQUES DEL TRABAJO
# ═══════════════════════════════════════════════════════════════════════════════
h1('Enfoques del trabajo')

body(
    'El enfoque metodológico del proyecto es ágil, basado en Scrum, con ciclo de '
    'vida adaptativo organizado en sprints de dos semanas. Esta elección permite '
    'focalizar el trabajo en los entregables de cada fase, iterar rápidamente sobre '
    'los resultados de la generación de contenido y adaptar el diseño de los agentes '
    'en función de la retroalimentación de los docentes revisores.'
)
body(
    'Las ceremonias que se realizarán son: Sprint Planning, para planificar y '
    'clarificar las tareas del próximo sprint; Sprint Review, para presentar el '
    'progreso al asesor académico; y Sprint Retrospective, para identificar '
    'oportunidades de mejora. Adicionalmente, se mantendrá un registro de avances '
    'en el repositorio de código (GitHub) para garantizar la trazabilidad del '
    'desarrollo.'
)

doc.add_paragraph()
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# RIESGOS Y MITIGACIÓN
# ═══════════════════════════════════════════════════════════════════════════════
h1('Riesgos y Mitigación')

body(
    'Los riesgos son eventos que, de materializarse, podrían impactar negativamente '
    'en el logro de los objetivos del proyecto. A continuación se presentan los '
    'principales riesgos identificados y sus estrategias de mitigación.'
)

risk_tbl = doc.add_table(rows=8, cols=5)
risk_tbl.style = 'Table Grid'
header_row(risk_tbl, ['#', 'Riesgo', 'Probabilidad', 'Impacto', 'Estrategia de mitigación'])
risks = [
    ('1',
     'Los LLMs generan contenido con baja alineación pedagógica o inconsistencias con los logros del sílabo.',
     'Media', 'Alto',
     'Mitigar: Implementar validaciones automáticas (TimeRevisor) y revisión obligatoria por docente revisor antes de aceptar el contenido como válido.'),
    ('2',
     'Cambios en las APIs de OpenAI o Google (precios, disponibilidad o modelos) afectan la implementación.',
     'Baja', 'Alto',
     'Evitar: Diseñar la arquitectura con abstracción de proveedores LLM para facilitar el reemplazo de modelos sin refactorización mayor.'),
    ('3',
     'Los docentes revisores no están disponibles para evaluar el contenido generado en los plazos del proyecto.',
     'Media', 'Alto',
     'Mitigar: Definir un calendario de revisión con al menos dos docentes revisores desde el inicio del proyecto.'),
    ('4',
     'La calidad del contenido generado varía significativamente entre tipos de cursos (técnicos vs. humanísticos).',
     'Media', 'Medio',
     'Mitigar: Incluir cursos de diferentes áreas en las pruebas de validación y ajustar los prompts por tipo de curso.'),
    ('5',
     'Incumplimiento de los plazos del ciclo académico de titulación UPC.',
     'Baja', 'Alto',
     'Evitar: Mantener un cronograma detallado con holgura en cada sprint y comunicar avances al asesor con anticipación.'),
    ('6',
     'Costo de consumo de APIs LLM superior al estimado por curso.',
     'Baja', 'Medio',
     'Mitigar: Monitorear el costo por ejecución durante las pruebas y ajustar el modelo o estrategia de prompts si el costo supera el límite definido.'),
    ('7',
     'Los documentos de entrada (sílabos, actas) tienen formatos inconsistentes que dificultan el parseo.',
     'Alta', 'Medio',
     'Mitigar: Implementar parsers robustos con manejo de excepciones y definir un conjunto de formatos válidos de entrada para las pruebas de validación.'),
]
for i, (num, risk, prob, impact, strategy) in enumerate(risks):
    row = risk_tbl.rows[i+1]
    row.cells[0].text = num
    row.cells[1].text = risk
    row.cells[2].text = prob
    row.cells[3].text = impact
    row.cells[4].text = strategy

doc.add_paragraph()
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# COSTOS DEL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
h1('Costos del proyecto')

body(
    'La solución tecnológica utiliza modelos de lenguaje de gran escala (LLMs) a '
    'través de APIs de pago por uso. A continuación se detallan los precios de los '
    'modelos utilizados, el costo por componente del proceso y el costo total '
    'estimado por curso generado, con base en datos reales de ejecución del sistema.'
)

body('Precios de los modelos LLM utilizados', bold=True)
llm_tbl = doc.add_table(rows=3, cols=3)
llm_tbl.style = 'Table Grid'
header_row(llm_tbl, ['Modelo', 'Precio entrada (por 1M tokens)', 'Precio salida (por 1M tokens)'])
llm_data = [
    ('Gemini 2.5-flash-lite (Google)', '$0.075 USD', '$0.30 USD'),
    ('GPT-4.1-mini (OpenAI)', '$0.40 USD', '$1.60 USD'),
]
for i, row_data in enumerate(llm_data):
    row = llm_tbl.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val

doc.add_paragraph()

body('Costo por componente del proceso (por curso — 18 semanas)', bold=True)
comp_tbl = doc.add_table(rows=12, cols=3)
comp_tbl.style = 'Table Grid'
header_row(comp_tbl, ['Componente / Agente', 'Fase', 'Costo estimado (USD)'])
comp_data = [
    ('KickOffAgent — parseo del acta de inicio', 'Diseño', '$0.0004'),
    ('CourseSchemaAgent — contexto del curso', 'Diseño', '$0.0001'),
    ('UniteSchemaSubAgent — esquema por unidad', 'Diseño', '$0.0037'),
    ('SyllabusActivityClassifier — clasificación', 'Diseño', '$0.0017'),
    ('UniteSchemaSubAgent — actividades', 'Diseño', '$0.0019'),
    ('IpesAgent — one_step_schema_recursos (18 sem.)', 'Generación', '$0.0045'),
    ('IpesAgent — one_step_introduccion (18 sem.)', 'Generación', '$0.0043'),
    ('IpesAgent — one_step_presentaciones (18 sem.)', 'Generación', '$0.0053'),
    ('IpesAgent — one_step_ejercicios (18 sem.)', 'Generación', '$0.0027'),
    ('CheckIpesAgent — QA del contenido', 'Aseguramiento', '$0.0002'),
    ('TOTAL IPES (sin material base)', '', '~$0.025 USD'),
]
for i, row_data in enumerate(comp_data):
    row = comp_tbl.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        if row_data[0].startswith('TOTAL'):
            row.cells[j].paragraphs[0].runs[0].bold = True if row.cells[j].paragraphs[0].runs else False

doc.add_paragraph()

body('Comparativa de costo: proceso manual vs. solución IA', bold=True)
cmp_tbl = doc.add_table(rows=3, cols=3)
cmp_tbl.style = 'Table Grid'
header_row(cmp_tbl, ['Proceso', 'Tiempo / Costo estimado', 'Observaciones'])
cmp_data = [
    ('Manual (docente virtualizador)',
     '~2 a 3 semanas × tarifa por hora',
     'Costo variable según volumen; no escalable sin incremento proporcional de horas.'),
    ('Solución IA (Gemini 2.5-flash-lite)',
     '~$0.025 USD por curso (IPES)\n~$0.064 USD por curso (flujo completo)',
     'Costo fijo por ejecución, independiente del número de cursos procesados en paralelo.'),
]
for i, row_data in enumerate(cmp_data):
    row = cmp_tbl.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val

body(
    'Nota: Los costos presentados se basan en datos reales obtenidos del registro de '
    'tokens del sistema (token_usage.jsonl), correspondientes a 566 llamadas de '
    'ejecución. El costo de $0.025 USD corresponde únicamente a la generación de '
    'materiales IPES (Introducción, Presentaciones y Ejercicios). El costo de $0.064 '
    'USD incluye adicionalmente la generación del material base de referencia '
    '(genera_material_base), que es opcional en el flujo de producción.'
)

# ── Costos operativos ─────────────────────────────────────────────────────────
doc.add_paragraph()
body('Costos operativos del proyecto', bold=True)
body(
    'Además del costo de consumo de APIs de IA, el proyecto involucra un equipo de '
    'trabajo con cinco roles activos bajo una metodología ágil (Scrum). A continuación '
    'se presenta el costo estimado por rol en función de las horas de participación '
    'proyectadas durante el ciclo de vida del proyecto (mayo–octubre 2026). '
    'Las tarifas por hora se derivan de las remuneraciones de referencia del mercado '
    'para cada perfil, calculadas sobre una base de 160 horas mensuales.'
)

op_tbl = doc.add_table(rows=7, cols=4)
op_tbl.style = 'Table Grid'
header_row(op_tbl, ['Rol', 'Horas estimadas', 'Tarifa/hora (S/.)', 'Costo total (S/.)'])
op_data = [
    ('Desarrollador IA',                              '400 hrs', 'S/. 31.25', 'S/. 12,500'),
    ('Project Manager',                               '120 hrs', 'S/. 37.50', 'S/. 4,500'),
    ('Tech Leader',                                   '100 hrs', 'S/. 37.50', 'S/. 3,750'),
    ('Product Owner',                                  '80 hrs', 'S/. 28.13', 'S/. 2,250'),
    ('Stakeholder / Sponsor\n(Jefe del equipo Modelos Emergentes)', '20 hrs', 'S/. 62.50', 'S/. 1,250'),
    ('TOTAL',                                         '720 hrs', '',          'S/. 24,250'),
]
for i, row_data in enumerate(op_data):
    row = op_tbl.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
    if row_data[0] == 'TOTAL':
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
body('Hardware asignado al equipo', bold=True)
hw2_tbl = doc.add_table(rows=2, cols=3)
hw2_tbl.style = 'Table Grid'
header_row(hw2_tbl, ['Recurso', 'Descripción', 'Cantidad'])
hw2_tbl.rows[1].cells[0].text = 'HP EliteBook Pro'
hw2_tbl.rows[1].cells[1].text = 'Laptop asignada por la institución al equipo de desarrollo del proyecto.'
hw2_tbl.rows[1].cells[2].text = '1 unidad'

doc.add_paragraph()
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CRONOGRAMA
# ═══════════════════════════════════════════════════════════════════════════════
h1('Cronograma del proyecto')

body(
    'El proyecto se ejecutará entre mayo y septiembre de 2026, organizado en cuatro '
    'fases alineadas a los objetivos específicos del trabajo de suficiencia profesional. '
    'Cada fase se desarrolla en sprints de dos semanas siguiendo la metodología Scrum.'
)

cron_tbl = doc.add_table(rows=17, cols=5)
cron_tbl.style = 'Table Grid'
header_row(cron_tbl, ['Sprint', 'Período', 'Fase / OE', 'Actividades principales', 'Entregable'])
cron_data = [
    # OE1 — Análisis
    ('Sprint 1', 'May 05 – May 18',  'OE1 — Análisis',
     'Revisión bibliográfica de LLMs en educación.\nAnálisis de tecnologías: GPT-4.1-mini, Gemini 2.5-flash-lite, LangGraph.',
     'Informe de análisis de tecnologías (borrador)'),
    ('Sprint 2', 'May 19 – Jun 01',  'OE1 — Análisis',
     'Análisis del proceso actual de virtualización.\nDocumentación del problema y antecedentes.',
     'Informe de análisis aprobado (IE1)'),
    # OE2 — Diseño
    ('Sprint 3', 'Jun 02 – Jun 15',  'OE2 — Diseño',
     'Diseño de la arquitectura multi-agente.\nDefinición de fases: Diseño, Aseguramiento, Generación.',
     'Diagrama de arquitectura (borrador)'),
    ('Sprint 4', 'Jun 16 – Jun 29',  'OE2 — Diseño',
     'Diseño detallado de agentes: KickOffAgent, CourseSchemaAgent, IpesAgent, CheckIpesAgent.\nDiagramas de flujo del proceso.',
     'Arquitectura aprobada (IE2)'),
    # OE3 — Implementación y Validación
    ('Sprint 5', 'Jun 30 – Jul 13',  'OE3 — Implementación',
     'Implementación de KickOffAgent y SyllabusActivityClassifier.\nPruebas unitarias de parseo de documentos de entrada.',
     'Agentes de diseño funcionales'),
    ('Sprint 6', 'Jul 14 – Jul 27',  'OE3 — Implementación',
     'Implementación de CourseSchemaAgent y UniteSchemaSubAgent.\nPruebas con sílabos reales.',
     'Fase Diseño implementada'),
    ('Sprint 7', 'Jul 28 – Aug 10',  'OE3 — Implementación',
     'Implementación de IpesAgent (Introducción, Presentaciones, Ejercicios).\nImplementación de CheckIpesAgent y TimeRevisor.',
     'Fase Generación implementada'),
    ('Sprint 8', 'Aug 11 – Aug 24',  'OE3 — Validación',
     'Ejecución del sistema con cursos de prueba.\nEvaluación del contenido por docentes revisores mediante rúbrica de calidad.',
     'Resultados de validación (borrador)'),
    ('Sprint 9', 'Aug 25 – Sep 07',  'OE3 — Validación',
     'Análisis comparativo: tiempo y costo IA vs. proceso manual.\nAjustes finales a los agentes según retroalimentación de revisores.',
     'Validación aprobada (IE3)'),
    # OE4 — Plan de continuidad
    ('Sprint 10', 'Sep 08 – Sep 21', 'OE4 — Continuidad',
     'Elaboración del plan de continuidad: roles de soporte, niveles de servicio, gestión de riesgos.',
     'Plan de continuidad (borrador)'),
    ('Sprint 11', 'Sep 22 – Sep 30', 'OE4 — Continuidad',
     'Revisión final del plan de continuidad.\nCierre de documentación del trabajo de suficiencia.',
     'Plan de continuidad aprobado (IE4)'),
    # Hito final
    ('Cierre', 'Octubre 2026',       'Entrega final',
     'Consolidación del informe final.\nPresentación ante el asesor académico.',
     'Trabajo de Suficiencia Profesional entregado'),
]
for i, (sprint, periodo, fase, actividades, entregable) in enumerate(cron_data):
    row = cron_tbl.rows[i+1]
    row.cells[0].text = sprint
    row.cells[1].text = periodo
    row.cells[2].text = fase
    row.cells[3].text = actividades
    row.cells[4].text = entregable

doc.add_paragraph()

body(
    'Nota: Las fechas son estimadas y podrán ajustarse según el calendario académico '
    'del Programa de Titulación UPC 2026. Los sprints de validación (OE3) están '
    'sujetos a la disponibilidad de los docentes revisores.'
)

doc.add_paragraph()
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# BIBLIOGRAFÍA
# ═══════════════════════════════════════════════════════════════════════════════
h1('Bibliografía')

refs = [
    'Hernández, R., Fernández, C., & Baptista, P. (2021). Metodología de la investigación (6.ª ed.). McGraw-Hill.',
    'Kasneci, E., Sessler, K., Küchemann, S., Bannert, M., Dementieva, D., Fischer, F., ... & Kasneci, G. (2023). ChatGPT for good? On opportunities and challenges of large language models for education. Learning and Individual Differences, 103, 102274.',
    'Tlili, A., Shehata, B., Adarkwah, M. A., Bozkurt, A., Hickey, D. T., Huang, R., & Agyemang, B. (2023). What if the devil is my guardian angel: ChatGPT as a case study of using chatbots in education. Smart Learning Environments, 10(1), 15.',
    'UNESCO. (2021). Education: From disruption to recovery. UNESCO. https://en.unesco.org/covid19/educationresponse',
    'Biggs, J. (1996). Enhancing teaching through constructive alignment. Higher Education, 32(3), 347-364.',
    'Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257-285.',
]
for ref in refs:
    body(ref)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# APROBACIÓN
# ═══════════════════════════════════════════════════════════════════════════════
h1('Aprobación')

aprov_tbl = doc.add_table(rows=3, cols=4)
aprov_tbl.style = 'Table Grid'
header_row(aprov_tbl, ['Nombre', 'Cargo', 'Firma', 'Fecha'])
aprov_data = [
    ('[Nombre Apellido]', 'Project Manager / Investigador Principal', '', 'Abril 2026'),
    ('Daniel Burga', 'Asesor Académico — UPC', '', 'Abril 2026'),
]
for i, (name, cargo, firma, fecha) in enumerate(aprov_data):
    row = aprov_tbl.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = cargo
    row.cells[2].text = firma
    row.cells[3].text = fecha

# ── Save ──────────────────────────────────────────────────────────────────────
out = 'D:/Proyectos actuales/proyecto_tesis_upc/Project_Charter_Tesis.docx'
doc.save(out)
print(f'Saved: {out}')
