from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_header_row(table, cols, color='1F3864'):
    row = table.rows[0]
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.text = col
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, color)

# ═══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 1
# ═══════════════════════════════════════════════════════════════════════════════

h1('CAPÍTULO 1: DEFINICIÓN DEL PROYECTO')

body(
    'En el presente capítulo se expone el contexto que motivó el desarrollo del proyecto. '
    'Se presentan los antecedentes de la problemática, una descripción de la organización '
    'donde se llevó a cabo la validación, el análisis del problema identificado, los '
    'objetivos propuestos, los indicadores de éxito y la descripción general del proyecto.'
)

# ── 1.1 Antecedentes ──────────────────────────────────────────────────────────

h2('Antecedentes')

body(
    'La pandemia de COVID-19 marcó un punto de inflexión en la educación superior del '
    'Perú. A partir del año 2020, y con mayor impacto durante 2021, las universidades '
    'se vieron obligadas a migrar masivamente hacia la modalidad virtual, sin contar con '
    'los procesos ni la infraestructura suficiente para sostener esa transición a escala. '
    'Según la UNESCO (2021), más del 90% de los sistemas educativos en América Latina '
    'implementaron alguna forma de educación a distancia durante la pandemia, lo que '
    'generó una demanda sin precedentes de contenido digital instruccional en plazos '
    'extremadamente cortos.'
)

body(
    'En este contexto, los equipos responsables del diseño y producción de cursos '
    'virtuales enfrentaron una doble presión: por un lado, la imposibilidad de coordinar '
    'presencialmente las etapas del proceso, y por otro, el incremento acelerado en el '
    'número de cursos que debían virtualizarse. El proceso de producción, que típicamente '
    'involucra a múltiples roles trabajando en cadena —diseñadores instruccionales, '
    'docentes revisores y virtualizadores—, se volvió especialmente vulnerable al '
    'trabajo remoto: cada entrega dependía de la disponibilidad del rol anterior, las '
    'reuniones de coordinación se dificultaron, y las versiones del contenido se '
    'multiplicaban sin control claro, generando reprocesos y retrasos acumulados.'
)

body(
    'Hernández et al. (2021) documentaron que, en instituciones de educación superior '
    'latinoamericanas, el principal obstáculo para la producción de contenido virtual '
    'durante la pandemia no fue tecnológico sino operativo: la coordinación entre los '
    'roles del equipo de diseño instruccional se deterioró significativamente al pasar '
    'al trabajo remoto, duplicando en muchos casos los tiempos de producción esperados. '
    'Este fenómeno fue particularmente agudo en la fase de redacción y revisión de '
    'materiales, donde la retroalimentación entre el diseñador y el experto temático '
    'requería múltiples iteraciones que, en entornos presenciales, se resolvían en '
    'reuniones cortas y directas.'
)

body(
    'Paralelamente, la generación del contenido instruccional —es decir, la redacción '
    'de los guiones de los recursos de video, las introducciones semanales y las '
    'consignas de actividades— representa históricamente la última etapa del proceso '
    'de producción, y por tanto la que concentra la mayor presión de entrega. En '
    'instituciones con alta demanda de cursos virtuales, esta etapa debe completarse '
    'en aproximadamente una semana por curso, lo que obliga al equipo a trabajar en '
    'paralelo con múltiples cursos simultáneamente, aumentando el riesgo de errores, '
    'inconsistencias pedagógicas y variabilidad en la calidad del contenido entre '
    'cursos de una misma institución.'
)

body(
    'Frente a este escenario, la inteligencia artificial —y en particular los modelos '
    'de lenguaje de gran escala (LLMs)— han emergido como una alternativa viable para '
    'asistir en la generación de contenido instruccional. Tlili et al. (2023) realizaron '
    'una revisión sistemática sobre el uso de IA generativa en educación, identificando '
    'que la generación de materiales instruccionales es una de las aplicaciones con '
    'mayor potencial de reducción de carga operativa, siempre que se mantengan '
    'mecanismos de revisión por expertos para garantizar la calidad pedagógica. '
    'En la misma línea, Kasneci et al. (2023) destacaron que los LLMs pueden generar '
    'contenido educativo coherente y alineado a objetivos de aprendizaje cuando se '
    'les proporcionan estructuras pedagógicas claras como punto de partida, lo que '
    'corresponde exactamente al enfoque adoptado en el presente proyecto: los agentes '
    'parten del sílabo y el acta de reunión de inicio para garantizar que todo el '
    'contenido generado esté anclado al diseño instruccional del curso.'
)

# ── 1.2 Descripción de la organización ───────────────────────────────────────

h2('Descripción de la organización')

body(
    'La organización de estudio es una institución de educación superior privada con '
    'presencia a nivel nacional en el Perú, que ofrece programas de pregrado y posgrado '
    'tanto en modalidad presencial como virtual. Cuenta con una amplia oferta académica '
    'orientada a diversas carreras profesionales y ha apostado por el crecimiento de su '
    'modalidad virtual como parte de su estrategia institucional.'
)

body(
    'En el ámbito de la educación virtual, la institución cuenta con un proceso '
    'estructurado de producción de cursos denominado virtualización, el cual comprende '
    'la definición del diseño instruccional, la redacción del contenido textual por '
    'semana académica y la posterior producción de los recursos multimedia. Dentro de '
    'este proceso, el docente virtualizador es el profesional encargado de redactar los '
    'materiales instruccionales denominados IPES (Introducción, Presentaciones y '
    'Ejercicios), que constituyen el contenido textual base de cada semana del curso.'
)

body(
    'El proceso de producción de un curso completo de 18 semanas toma en promedio dos '
    'semanas de trabajo, durante las cuales el docente virtualizador genera los guiones '
    'de narración para los recursos de video, las introducciones a cada semana y las '
    'consignas de los ejercicios y actividades, tanto calificadas como no calificadas. '
    'La institución establece criterios pedagógicos precisos para este contenido, '
    'incluyendo la alineación con los logros de aprendizaje del sílabo, el respeto de '
    'restricciones de tiempo de estudio y la selección de tipos de recursos según el '
    'nivel taxonómico del logro. Una vez generado el contenido, un docente revisor '
    'valida su calidad y consistencia antes de pasar a la fase de producción multimedia.'
)

# ── 1.3 Análisis del problema ─────────────────────────────────────────────────

h2('Análisis del problema')

body(
    'La demora en la gestión de contenido y materiales de los cursos virtuales en la '
    'institución constituye el problema central identificado. La producción de material '
    'educativo para los cursos virtuales enfrenta un cuello de botella operativo: la '
    'redacción de los materiales instruccionales de cada semana académica recae sobre '
    'el docente virtualizador, un profesional especializado que escribe manualmente las '
    'introducciones, guiones de narración para videos de presentación y consignas de '
    'actividades, asegurando que cada pieza esté alineada a los logros de aprendizaje '
    'del curso. Esta situación se origina en las siguientes causas:'
)

causes = [
    (
        '1. Brecha operativa en la generación de contenido',
        'El proceso depende exclusivamente de la disponibilidad y capacidad del docente '
        'virtualizador, lo que hace que la producción de un curso completo tome en promedio '
        'dos semanas. Esto limita la capacidad institucional de escalar la oferta académica '
        'virtual frente a una demanda creciente.'
    ),
    (
        '2. Costo operativo elevado',
        'Al remunerar al docente virtualizador por horas trabajadas, el gasto de producción '
        'aumenta proporcionalmente con la cantidad de cursos, sin que exista actualmente un '
        'mecanismo que permita reducir esa carga sin comprometer la calidad del contenido.'
    ),
    (
        '3. Baja integración de inteligencia artificial en los procesos de virtualización',
        'La incorporación de IA en procesos educativos es una tendencia que las instituciones '
        'de educación superior deberán atender progresivamente. Si la generación de contenido '
        'instruccional no se integra a este paradigma, las universidades perderán la '
        'oportunidad de conectar este proceso con otras etapas del flujo de virtualización '
        'que en el futuro también podrían ser asistidas por IA.'
    ),
    (
        '4. Riesgo operativo a largo plazo',
        'De no abordarse esta problemática, la generación de contenido continuará siendo '
        'un proceso manual, costoso y poco escalable, representando una limitación '
        'estructural para las instituciones que buscan ampliar su oferta virtual de '
        'manera sostenible.'
    ),
]

for title, detail in causes:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(detail)
    p.paragraph_format.space_after = Pt(4)

body('A continuación, se sintetiza la problemática del proyecto y sus causas subyacentes.')

# Problem / Causes table
tbl = doc.add_table(rows=2, cols=2)
tbl.style = 'Table Grid'
add_header_row(tbl, ['Problema', 'Causas'])

r = tbl.rows[1]
r.cells[0].text = (
    'Demora en la producción de material educativo para los cursos virtuales, '
    'derivada de la dependencia del proceso en el docente virtualizador y la ausencia '
    'de herramientas tecnológicas que optimicen la generación de contenido instruccional.'
)
r.cells[1].text = (
    '• Dependencia exclusiva del docente virtualizador para producir el contenido semanal.\n'
    '• Remuneración por hora que escala linealmente con el volumen de cursos.\n'
    '• Ausencia de herramientas de IA en la fase de generación de contenido instruccional.\n'
    '• Proceso no escalable que limita el crecimiento sostenible de la oferta virtual.'
)

doc.add_paragraph()

# ── 1.4 Objetivos ─────────────────────────────────────────────────────────────

h2('Objetivos')

body(
    'A continuación se describen los objetivos propuestos para dar solución a la '
    'problemática identificada. Se plantea el objetivo general que define la meta '
    'central del proyecto, seguido de los objetivos específicos que detallan las '
    'acciones concretas para alcanzarlo.'
)

h3('Objetivo general')

body(
    'Implementar la solución tecnológica basada en inteligencia artificial para '
    'optimizar el tiempo, la generación de contenido y los cursos virtuales en '
    'la universidad.'
)

h3('Objetivos específicos')

oes = [
    ('OE1',
     'Analizar la solución tecnológica basada en IA para optimizar el tiempo en la '
     'generación de contenido de cursos virtuales.'),
    ('OE2',
     'Diseñar la arquitectura de la solución tecnológica basada en IA para optimizar '
     'el tiempo en la generación de contenido de cursos virtuales en universidades del Perú.'),
    ('OE3',
     'Implementar y validar la solución tecnológica basada en IA para optimizar el '
     'tiempo en la generación de contenido de cursos virtuales mediante métricas de rendimiento.'),
    ('OE4',
     'Elaborar un plan de continuidad que asegure la viabilidad técnica de la solución '
     'tecnológica propuesta en el tiempo.'),
]

for code, text in oes:
    p = doc.add_paragraph()
    p.add_run(f'{code}: ').bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ── 1.5 Indicadores de éxito ──────────────────────────────────────────────────

h2('Indicadores de éxito')

body(
    'En la siguiente tabla se presentan los indicadores de éxito definidos para el '
    'proyecto, asociados a cada objetivo específico.'
)

ie_tbl = doc.add_table(rows=5, cols=3)
ie_tbl.style = 'Table Grid'
add_header_row(ie_tbl, ['Indicador', 'Descripción', 'Objetivo'])

ie_data = [
    ('IE1',
     'Aprobación por el asesor académico del análisis de tecnologías, herramientas '
     'y enfoques de IA para la generación de contenido instruccional.',
     'OE1'),
    ('IE2',
     'Aprobación por el asesor académico de la arquitectura de agentes diseñada para '
     'las fases de diseño, aseguramiento y generación de contenido.',
     'OE2'),
    ('IE3',
     'Aprobación por el asesor académico de la implementación y validación de la '
     'solución, con resultados favorables en las métricas de rendimiento establecidas.',
     'OE3'),
    ('IE4',
     'Aprobación por el asesor académico del plan de continuidad que garantice la '
     'viabilidad técnica y operativa de la solución.',
     'OE4'),
]

for i, (code, desc, obj) in enumerate(ie_data):
    row = ie_tbl.rows[i + 1]
    row.cells[0].text = code
    row.cells[1].text = desc
    row.cells[2].text = obj

doc.add_paragraph()

# ── 1.6 Descripción del proyecto ──────────────────────────────────────────────

h2('Descripción del proyecto')

body(
    'El proyecto propone el uso de agentes de inteligencia artificial para optimizar '
    'la generación del contenido textual de cursos virtuales, abordando específicamente '
    'la fase donde el docente virtualizador actualmente trabaja de forma manual. El '
    'contenido objetivo son los materiales denominados IPES (Introducción, Presentaciones '
    'y Ejercicios), que representan el núcleo textual de cada semana académica de un '
    'curso virtual.'
)

body(
    'El proceso de generación se estructura en tres fases. En primer lugar, la fase de '
    'diseño del curso, donde los agentes procesan los documentos de entrada —el sílabo, '
    'el acta de reunión inicial y la bibliografía— para construir la estructura del '
    'contenido: temas, subtemas y apartados de cada semana, alineados a los logros de '
    'aprendizaje definidos institucionalmente. En segundo lugar, la fase de '
    'aseguramiento, en la que los agentes verifican que la estructura generada sea '
    'consistente con los logros del curso y cumpla con las restricciones pedagógicas '
    'establecidas. Finalmente, la fase de generación de contenido, donde los agentes '
    'producen el material textual de cada semana: la introducción, los guiones de '
    'narración para los recursos de presentación y las consignas de los ejercicios, '
    'tanto calificados como no calificados, dando como resultado un documento Word por '
    'semana durante las 18 semanas que comprende el curso.'
)

body(
    'La efectividad del contenido generado será evaluada por docentes revisores, '
    'profesionales que en el proceso actual validan el trabajo producido por los '
    'docentes virtualizadores, por lo que cuentan con el criterio necesario para '
    'juzgar la calidad, coherencia pedagógica y alineación del contenido generado '
    'por los agentes de IA.'
)

out = 'D:/Proyectos actuales/proyecto_tesis_upc/Capitulo1_Tesis_v2.docx'
doc.save(out)
print(f'Saved: {out}')
