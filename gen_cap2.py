# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def h1(text):
    return doc.add_heading(text, level=1)

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def h4(text):
    return doc.add_heading(text, level=4)

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def body_bold_start(bold_text, rest_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    p.add_run(rest_text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ': ')
        r.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def ref_note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_header_row(table, cols, color='1F3864'):
    row = table.rows[0]
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.text = col
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, color)

def spacer():
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 2 — MARCO TEÓRICO
# ═══════════════════════════════════════════════════════════════════════════════

h1('CAPÍTULO 2: MARCO TEÓRICO')

# ═══════════════════════════════════════════════════════════════════════════════
# 2.1  PROCESO DE PRODUCCIÓN INSTRUCCIONAL
# ═══════════════════════════════════════════════════════════════════════════════

h2('2.1  El Proceso de Producción Instruccional')

body_bold_start(
    'En esta sección se describen los elementos que componen el proceso de producción '
    'de cursos virtuales en una institución de educación superior: los roles del equipo, '
    'los documentos de entrada que definen el curso',
    ', los artefactos que se producen y el catálogo de tipos de recursos educativos.'
)

# ── 2.1.1 Roles ───────────────────────────────────────────────────────────────

h3('2.1.1  Roles del equipo de producción')

body(
    'Se explicará cada uno de los perfiles que intervienen en el proceso de '
    'virtualización de cursos, su responsabilidad específica dentro del flujo '
    'de producción y la relación entre ellos. Los roles a desarrollar son:'
)

roles = [
    ('DDA (Dirección de Diseño Académico)',
     'Rol con visión macro del proceso de producción. '
     'Puede intervenir y tomar decisiones sobre cualquier curso en cualquier momento. '
     'No necesariamente es experto en el contenido de cada asignatura.'),
    ('LxD (Learning Designer / Diseñador Instruccional)',
     'Profesional que diseña la estructura instruccional del curso. '
     'Conoce cómo debe estructurarse un curso pero no es el experto en la materia. '
     'Es quien genera el IPES.'),
    ('Docente Virtualizador / Experto Temático',
     'Experto en el contenido de la asignatura y responsable de crear '
     'los materiales instruccionales. Es el rol que el sistema replica mediante IA.'),
    ('Docente Revisor',
     'Experto que revisa y valida los materiales generados —ya sea por humanos o por IA— '
     'antes de que pasen a producción. Es el validador final de calidad.'),
    ('Área de Evaluación',
     'Equipo especializado en la construcción de instrumentos de evaluación (consignas, '
     'rúbricas, listas de cotejo). Verifica que las actividades calificadas estén bien '
     'construidas pedagógicamente.'),
]

for title, desc in roles:
    bullet(desc, title)

spacer()

# ── 2.1.2 Inputs ──────────────────────────────────────────────────────────────

h3('2.1.2  Documentos de entrada: Sílabo y Kickoff')

body(
    'Se describirán los dos documentos que el sistema recibe como entrada y que '
    'concentran toda la información necesaria para producir el contenido de un curso:'
)

inputs_data = [
    ('Sílabo',
     'Documento oficial del curso (formato Excel) que contiene la estructura académica: '
     'unidades, semanas, temas, logros de aprendizaje y actividades calificadas. '
     'Es la fuente de verdad estructural del curso.'),
    ('Kickoff / Acta de Reunión de Inicio',
     'Documento que recoge los acuerdos tomados entre el equipo académico y el experto '
     'antes de iniciar la producción. Incluye: perfil del estudiante, enfoque de diseño '
     '(teórico / práctico / mixto), temas complejos, fórmula de evaluación y criterios '
     'específicos del curso. Contiene información que el sílabo no recoge y que '
     'condiciona decisiones de diseño.'),
]

for title, desc in inputs_data:
    bullet(desc, title)

spacer()

# ── 2.1.3 Outputs ─────────────────────────────────────────────────────────────

h3('2.1.3  Artefactos producidos: IPES, Material Base y Consignas')

body(
    'Se explicarán los tres artefactos principales que produce el proceso de '
    'virtualización, con énfasis en la relación secuencial entre ellos:'
)

outputs = [
    ('IPES (Introducción, Presentaciones y Ejercicios)',
     'Primer output del sistema. Documento de planificación instruccional por semana '
     'que define la situación de apertura, los recursos de presentación con su tipo y '
     'propósito, y el ejercicio de la semana (calificado o no calificado). '
     'El IPES no es el contenido final; es la hoja de ruta que indica qué producir.'),
    ('Material Base / Base Content',
     'Último output del proceso. Usando el IPES como entrada, se genera el contenido '
     'real que consumirán los estudiantes: guiones de narración para videos, texto de '
     'separatas y manuales, pantallas para recursos interactivos, instrucciones y '
     'resoluciones de ejercicios. Es el artefacto final de la cadena de producción.'),
    ('Consignas',
     'Instrumentos de evaluación formales: rúbricas, listas de cotejo y escalas de '
     'valoración. Definen cómo se evaluará el desempeño del estudiante en las '
     'actividades calificadas. Son revisados por el Área de Evaluación.'),
]

for title, desc in outputs:
    bullet(desc, title)

spacer()

# ── 2.1.4 Catálogo de recursos educativos ────────────────────────────────────

h3('2.1.4  Catálogo de recursos educativos')

body(
    'Se presentará el conjunto de tipos de recursos educativos que el sistema maneja. '
    'La clasificación de recursos es institucional: cada universidad puede definir '
    'su propio catálogo; el sistema utiliza el de la institución donde se validó el proyecto.'
)

res_tbl = doc.add_table(rows=14, cols=2)
res_tbl.style = 'Table Grid'
add_header_row(res_tbl, ['Tipo de Recurso', 'Categoría'])

res_data = [
    ('H5P',                    'Contenido interactivo'),
    ('HTML',                   'Contenido interactivo'),
    ('Rise',                   'Contenido interactivo'),
    ('Storyline',              'Contenido interactivo'),
    ('Video Explicativo',      'Audiovisual'),
    ('Video Demo',             'Audiovisual'),
    ('Video Interactivo',      'Audiovisual'),
    ('Podcast',                'Audiovisual'),
    ('Separata',               'Lectura'),
    ('Manual',                 'Lectura'),
    ('Lectura',                'Lectura'),
    ('Organizador Visual',     'Visual / gráfico'),
    ('Lectura Complementaria', 'Lectura opcional'),
]

for i, (tipo, cat) in enumerate(res_data):
    row = res_tbl.rows[i + 1]
    row.cells[0].text = tipo
    row.cells[1].text = cat

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 2.2  FUNDAMENTOS PEDAGÓGICOS
# ═══════════════════════════════════════════════════════════════════════════════

h2('2.2  Fundamentos Pedagógicos')

body(
    'En esta sección se presenta el marco teórico del aprendizaje que sustenta '
    'las reglas pedagógicas codificadas en el sistema.'
)

# ── 2.2.1 Taxonomía de Bloom ──────────────────────────────────────────────────

h3('2.2.1  Taxonomía de Bloom')

body(
    'La Taxonomía de Bloom, en su revisión de Anderson & Krathwohl (2001), constituye '
    'el marco central de clasificación de objetivos de aprendizaje sobre el que el sistema '
    'fundamenta sus decisiones de selección de recursos. Se explicará en detalle dado '
    'su rol estructural en la lógica del sistema. Los puntos a cubrir son:'
)

bloom_points = [
    'Los seis niveles cognitivos y sus verbos asociados '
    '(Recordar, Comprender, Aplicar, Analizar, Evaluar, Crear), según la revisión '
    'de Anderson & Krathwohl (2001) sobre la taxonomía original de Bloom et al. (1956).',

    'Cómo el verbo del logro de aprendizaje semanal determina el tipo de recurso '
    'a utilizar: la tabla verbo × densidad de contenido → tipo de recurso '
    'que el sistema aplica en cada semana del curso.',

    'La jerarquía de logros: Logro del Curso → Logro de Unidad → '
    'Logro de Semana → Propósito del Recurso, y por qué cada nivel '
    'debe ser consistente con el superior (principio de alineación vertical).',

    'La distinción entre contenido procedimental (requiere recurso demostrativo), '
    'teórico (requiere recurso expositivo) y básico (flexible), '
    'y cómo el verbo del logro permite identificar cada tipo.',
]

for pt in bloom_points:
    bullet(pt)

ref_note(
    'Referencias clave: Bloom, B. S., Engelhart, M. D., Furst, E. J., Hill, W. H., & Krathwohl, D. R. (1956). '
    'Taxonomy of Educational Objectives: The Classification of Educational Goals. David McKay. | '
    'Anderson, L. W., & Krathwohl, D. R. (2001). A Taxonomy for Learning, Teaching, and Assessing: '
    "A Revision of Bloom's Taxonomy of Educational Objectives. Longman."
)

spacer()

# ── 2.2.2 Rúbrica de evaluación de calidad del contenido generado ─────────────

h3('2.2.2  Rúbrica de evaluación de calidad del contenido generado')

body(
    'Se describirá el instrumento de evaluación construido para que los Docentes Revisores '
    'valoren la calidad pedagógica del contenido generado por el sistema. '
    'A diferencia de las métricas técnicas, este instrumento mide la calidad percibida '
    'por expertos a través de seis criterios que cubren alineación, exactitud, rigor, '
    'adecuación al nivel del alumno, calidad de ejemplos y capacidad motivadora del material.'
)

rub_tbl = doc.add_table(rows=7, cols=4)
rub_tbl.style = 'Table Grid'
add_header_row(rub_tbl, ['Criterio', 'Estándar Esperado (5 pts)', 'En proceso (3 pts)', 'Inicial (1 pt)'])

rubric_data = [
    (
        'Alineación del contenido',
        'El contenido tiene una conexión clara con el logro de la unidad y cubre el tema con un alcance completo, sin dejar vacíos de información.',
        'El contenido está relacionado con el logro de la unidad pero incluye material solo tangencialmente relacionado o que divaga.',
        'El contenido no está relacionado con sus logros y/o tiene material con un alcance insuficiente.',
    ),
    (
        'Exactitud del contenido',
        'El contenido del material es completamente factualmente correcto y tiene suficientes matices para el tema.',
        'El contenido del material tiene errores factuales superficiales como terminología equivocada, simplificaciones o imprecisiones, pero encaja con el tema.',
        'El contenido del material tiene errores factuales profundos o no encaja con el tema.',
    ),
    (
        'Rigor académico',
        'El contenido del material es profundo, creíble y cita / se basa en fuentes confiables.',
        'El contenido es superficial o hace afirmaciones cuestionables / poco creíbles.',
        'El contenido no tiene sustento o hace afirmaciones sin soporte.',
    ),
    (
        'Propiedad para el nivel del alumno',
        'El vocabulario, el nivel de dificultad y complejidad del material son los adecuados para el nivel del alumno.',
        'El vocabulario, nivel de dificultad y complejidad son ligeramente muy avanzados o muy simples para el nivel del alumno.',
        'El vocabulario, nivel de dificultad y complejidad no tienen conexión con el nivel del alumno.',
    ),
    (
        'Calidad de ejemplos',
        'El contenido tiene ejemplos relevantes y concretos que suplementan los conceptos explicados.',
        'El contenido tiene ejemplos genéricos o solo parcialmente relevantes a los conceptos explicados.',
        'El contenido tiene ejemplos confusos, pobres o no existentes.',
    ),
    (
        'Probabilidad de que el alumno complete el material',
        'El contenido puede generar interés intrínsecamente y haría que los alumnos lo revisen voluntariamente.',
        'El contenido se presta a que los alumnos lo lean si es requerido pero no es motivador.',
        'El contenido se presta a que los alumnos lo revisen superficialmente o lo salten aunque sea requerido.',
    ),
]

for i, (criterio, estandar, en_proceso, inicial) in enumerate(rubric_data):
    row = rub_tbl.rows[i + 1]
    row.cells[0].text = criterio
    row.cells[1].text = estandar
    row.cells[2].text = en_proceso
    row.cells[3].text = inicial

spacer()

body(
    'Se explicará cada criterio en relación con el tipo de contenido que el sistema '
    'genera, y por qué estos seis aspectos son los relevantes para juzgar la calidad '
    'de un material instruccional generado por IA en el contexto universitario.'
)

ref_note(
    'Referencias clave: Wiggins, G., & McTighe, J. (2005). Understanding by Design (2nd ed.). ASCD. | '
    'Kasneci, E., Sessler, K., Küchemann, S., Bannert, M., Dementieva, D., Fischer, F., & Kasneci, G. (2023). '
    'ChatGPT for good? On opportunities and challenges of large language models for education. '
    'Learning and Individual Differences, 103, 102274.'
)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 2.3  FUNDAMENTOS TÉCNICOS
# ═══════════════════════════════════════════════════════════════════════════════

h2('2.3  Fundamentos Técnicos')

body(
    'En esta sección se presentan los conceptos de inteligencia artificial e ingeniería '
    'de software sobre los que se construyó el sistema, organizados en dos capas: '
    'los componentes de la capa de agentes de IA y los componentes de la infraestructura backend.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2.3.1  AGENTES DE IA
# ═══════════════════════════════════════════════════════════════════════════════

h3('2.3.1  Agentes de IA')

body(
    'Se presentan los conceptos que conforman la capa de inteligencia artificial del sistema: '
    'el mecanismo de instrucción de los modelos, la estructura de sus salidas, '
    'la especialización por agente, el control de flujo y el framework de orquestación.'
)

# ── Prompt Engineering ────────────────────────────────────────────────────────

h4('Prompt Engineering')

body(
    'El diseño de prompts es la principal interfaz de programación con los modelos de lenguaje '
    'grandes. En este trabajo, cada agente del sistema está definido por un prompt que '
    'especifica su rol, sus reglas obligatorias y sus lineamientos de criterio. '
    'Se explicará por qué es relevante para esta tesis:'
)

bullet(
    'Qué es un prompt y la separación entre system prompt '
    '(instrucciones permanentes de rol y reglas) y user prompt '
    '(contexto variable de cada ejecución).',
)
bullet(
    'El patrón de hard rules vs. soft guidelines: reglas no negociables '
    'que el agente debe respetar siempre, versus lineamientos donde '
    'aplica criterio propio. Este patrón gobierna todos los agentes del sistema.',
)
bullet(
    'El proceso iterativo de refinamiento de prompts con validación por expertos '
    'como metodología de desarrollo de la capa de IA (Wei et al., 2022).',
)

ref_note(
    'Referencias: Brown, T. B., et al. (2020). Language models are few-shot learners. NeurIPS, 33, 1877-1901. | '
    'Liu, P., et al. (2023). Pre-train, prompt, and predict. ACM Computing Surveys, 55(9), 1-35. | '
    'Wei, J., et al. (2022). Chain-of-thought prompting elicits reasoning in large language models. NeurIPS, 35.'
)
spacer()

# ── Structured Outputs ────────────────────────────────────────────────────────

h4('Structured Outputs (Salidas Estructuradas)')

body(
    'Para que la solución funcione como pipeline encadenado, cada agente debe producir '
    'datos en formato JSON validado, no texto libre. '
    'Se explicará por qué es fundamental en esta tesis:'
)

bullet(
    'JSON mode / structured output: el modelo es forzado a producir '
    'un JSON que cumple un esquema predefinido (OpenAI, 2024; Google DeepMind, 2024). '
    'Permite que el output de un agente sea el input directo del siguiente.',
)
bullet(
    'Pydantic como contrato de datos: cada agente define su modelo de '
    'entrada y salida mediante clases Pydantic, garantizando tipos y '
    'campos correctos en toda la cadena de producción.',
)

ref_note(
    'Referencias: OpenAI. (2024). Structured outputs. OpenAI Platform Documentation. '
    'https://platform.openai.com/docs/guides/structured-outputs'
)
spacer()

# ── Sistemas Multi-Agente ─────────────────────────────────────────────────────

h4('Sistemas Multi-Agente')

body(
    'El sistema distribuye la generación de contenido entre múltiples agentes especializados '
    'en lugar de un único agente generalista. '
    'Se explicará por qué este enfoque es relevante para esta tesis:'
)

bullet(
    'Qué es un agente de IA: una instancia de un modelo de lenguaje '
    'con un rol, instrucciones y responsabilidad específica.',
)
bullet(
    'Por qué especializar: un agente con una sola responsabilidad bien definida '
    'produce resultados más precisos y confiables. En este sistema, '
    'agentes distintos manejan el parseo, el esquema curricular, '
    'la generación de contenido y la validación de calidad (Park et al., 2023).',
)

ref_note(
    'Referencias: Park, J. S., et al. (2023). Generative agents: Interactive simulacra of human behavior. UIST. | '
    'Wu, Q., et al. (2023). AutoGen: Enabling next-gen LLM applications via multi-agent conversation. arXiv:2308.08155.'
)
spacer()

# ── Nodos y Grafos de Estado ──────────────────────────────────────────────────

h4('Nodos y Grafos de Estado')

body(
    'La orquestación de agentes en el sistema se implementa como un grafo de estado dirigido, '
    'donde cada nodo representa un paso de procesamiento. '
    'Se explicará por qué es relevante para esta tesis:'
)

bullet(
    'Un nodo como unidad de procesamiento: puede invocar un LLM, '
    'ejecutar código determinístico o tomar una decisión condicional. '
    'El estado es compartido entre todos los nodos del grafo.',
)
bullet(
    'Aristas condicionales y ciclos de reintento: el flujo puede bifurcarse '
    'según el resultado de un nodo (validación pasa → continuar, '
    'validación falla → reintentar con feedback). '
    'Este mecanismo es el que garantiza el cumplimiento de restricciones pedagógicas.',
)
spacer()

# ── LangChain y LangGraph ─────────────────────────────────────────────────────

h4('Agent Framework: LangChain y LangGraph')

body(
    'LangChain y LangGraph son las herramientas de orquestación sobre las que '
    'se implementó el sistema. Se explicará por qué son relevantes para esta tesis:'
)

bullet(
    'LangChain: capa de abstracción sobre múltiples proveedores de LLM '
    '(OpenAI, Google Gemini) que estandariza la invocación de modelos, '
    'el manejo de prompts y el procesamiento de outputs.',
)
bullet(
    'LangGraph: extensión de LangChain que implementa grafos de estado dirigidos, '
    'habilitando flujos condicionales, ciclos de reintento y estado compartido entre nodos. '
    'Supera al prompt chaining simple al permitir lógica de control real sobre el flujo '
    'de generación (Chase, 2024).',
)

ref_note(
    'Referencias: Chase, H. (2024). LangGraph: Building stateful, multi-actor applications with LLMs. '
    'LangChain, Inc. https://langchain-ai.github.io/langgraph/ | '
    'Yao, S., et al. (2023). ReAct: Synergizing reasoning and acting in language models. ICLR.'
)
spacer()

# ── Python ────────────────────────────────────────────────────────────────────

h4('Python como ecosistema de desarrollo de IA')

body(
    'Python es el lenguaje de implementación del backend. '
    'Se justificará su uso en el contexto de esta tesis:'
)

bullet(
    'Ecosistema: concentra las bibliotecas más utilizadas para modelos de lenguaje '
    '(LangChain, LangGraph), procesamiento documental (pdfplumber, python-docx, pandas) '
    'y exposición de APIs (FastAPI).',
)
bullet(
    'FastAPI y Pydantic: FastAPI expone los agentes como endpoints HTTP; '
    'Pydantic valida los modelos de datos y hace cumplir los structured outputs '
    'en toda la cadena de agentes.',
)

ref_note(
    'Referencias: Raschka, S., Patterson, J., & Nolet, C. (2020). '
    'Machine learning in Python: Main developments and technology trends. Information, 11(4), 193.'
)
spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 2.3.2  INFRAESTRUCTURA BACKEND
# ═══════════════════════════════════════════════════════════════════════════════

h3('2.3.2  Infraestructura Backend')

body(
    'Se presentan los conceptos de arquitectura de software y servicios cloud '
    'que soportan la ejecución del sistema en producción. '
    'Dado que la generación de contenido con agentes LLM es un proceso de '
    'varios minutos, el backend requiere un modelo asíncrono basado en colas '
    'y cómputo serverless.'
)

# ── Procesamiento Asíncrono ───────────────────────────────────────────────────

h4('Procesamiento Asíncrono')

body(
    'El procesamiento asíncrono es el patrón arquitectónico que permite desacoplar '
    'la recepción de una solicitud de su ejecución. En un modelo síncrono, el cliente '
    'espera bloqueado hasta que el servidor responde; para procesos de varios minutos '
    '—como la generación completa de un curso— esto genera timeouts inaceptables '
    '(Kleppmann, 2017). En el modelo asíncrono adoptado, el cliente envía la solicitud '
    'y recibe inmediatamente un identificador de trabajo (job_id); el procesamiento '
    'ocurre en segundo plano y, al finalizar, el sistema notifica al cliente mediante webhook. '
    'Este patrón sigue el modelo productor-consumidor: el componente que recibe la solicitud '
    'deposita un mensaje en una cola y el worker lo toma y ejecuta el pipeline de agentes '
    'de forma independiente (Richardson, 2018).'
)

ref_note(
    'Referencias: Kleppmann, M. (2017). Designing Data-Intensive Applications. O\'Reilly Media. | '
    'Richardson, C. (2018). Microservices Patterns. Manning Publications.'
)
spacer()

# ── AWS ───────────────────────────────────────────────────────────────────────

h4('Amazon Web Services (AWS)')

body(
    'Amazon Web Services (AWS) es la plataforma cloud elegida para el despliegue del backend. '
    'Ofrece un ecosistema de servicios gestionados con modelo de pago por uso e integración '
    'nativa entre mensajería (SQS, SNS) y cómputo serverless (Lambda), lo que reduce '
    'la complejidad operativa frente a infraestructura autogestionada (Wittig & Wittig, 2018). '
    'En el sistema se utilizan cuatro servicios de AWS: API Gateway como punto de entrada HTTP, '
    'Lambda como motor de ejecución de los agentes, SQS como cola de trabajos y SNS '
    'como capa de notificaciones. Cada uno se describe a continuación.'
)

ref_note(
    'Referencias: Wittig, M., & Wittig, A. (2018). Amazon Web Services in Action (2nd ed.). Manning Publications. | '
    'Amazon Web Services. (2024). AWS Documentation. https://docs.aws.amazon.com'
)
spacer()

# ── SQS ───────────────────────────────────────────────────────────────────────

h4('Amazon SQS (Simple Queue Service)')

body(
    'Amazon SQS es el servicio de cola de mensajes completamente gestionado que administra '
    'los trabajos de generación de contenido del sistema. Cada solicitud de generación '
    'de curso se encola como un mensaje; SQS actúa como buffer ante múltiples solicitudes '
    'simultáneas y garantiza que ningún trabajo se pierda ante fallos del worker, '
    'dado que el mensaje permanece en la cola hasta ser procesado exitosamente. '
    'La integración nativa con Lambda permite que SQS dispare la función de procesamiento '
    'automáticamente cuando hay mensajes disponibles, sin necesidad de polling activo '
    'por parte del sistema.'
)

ref_note(
    'Referencias: Amazon Web Services. (2024). Amazon SQS Developer Guide. https://docs.aws.amazon.com/sqs/'
)
spacer()

# ── SNS ───────────────────────────────────────────────────────────────────────

h4('Amazon SNS (Simple Notification Service)')

body(
    'Amazon SNS es el servicio de mensajería de publicación/suscripción (pub/sub) '
    'que comunica la finalización de un trabajo de generación a los sistemas interesados. '
    'Cuando el worker concluye el procesamiento, publica un evento en un SNS topic; '
    'los suscriptores registrados —otros servicios o webhooks— reaccionan a ese evento '
    'sin que el worker necesite conocerlos directamente. '
    'Esta arquitectura permite aplicar el patrón fanout: SNS distribuye el evento '
    'de finalización a múltiples destinos en paralelo, desacoplando la lógica '
    'de notificación del código de generación.'
)

ref_note(
    'Referencias: Amazon Web Services. (2024). Amazon SNS Developer Guide. https://docs.aws.amazon.com/sns/'
)
spacer()

# ── Webhooks ──────────────────────────────────────────────────────────────────

h4('Webhooks')

body(
    'Un webhook es una llamada HTTP que el servidor realiza automáticamente a una URL '
    'predefinida por el cliente cuando un evento específico ocurre. A diferencia del polling '
    '—donde el cliente consulta repetidamente el estado del trabajo—, el webhook empuja '
    'la notificación en el momento exacto en que el resultado está listo, eliminando '
    'el tráfico innecesario durante el procesamiento. En el sistema, al enviar la solicitud '
    'de generación el cliente registra una URL de callback; cuando el worker finaliza, '
    'SNS dispara el webhook hacia esa URL con el resultado del trabajo '
    '(Fielding & Taylor, 2002).'
)

ref_note(
    'Referencias: Fielding, R. T., & Taylor, R. N. (2002). '
    'Principled design of the modern web architecture. ACM TOIT, 2(2), 115-150.'
)
spacer()

# ── Lambda ────────────────────────────────────────────────────────────────────

h4('AWS Lambda (Cómputo Serverless)')

body(
    'AWS Lambda es el servicio de cómputo serverless que ejecuta el pipeline de agentes LLM '
    'en el backend. El desarrollador despliega el código como una función y AWS gestiona '
    'el provisionamiento, el escalado y la ejecución en respuesta a eventos, sin necesidad '
    'de administrar servidores. Lambda admite un timeout máximo de 15 minutos por ejecución, '
    'límite que ninguna generación del sistema supera —ni siquiera el pipeline multi-agente '
    'completo—, lo que hace que el servicio cubra el caso de uso sin requerir instancias '
    'de larga ejecución. Además, Lambda escala horizontalmente de forma automática '
    'y cobra únicamente por el tiempo de cómputo consumido durante la ejecución, '
    'lo que reduce el costo operativo para un sistema de uso no continuo '
    '(Sbarski & Kroonenburg, 2017).'
)

ref_note(
    'Referencias: Amazon Web Services. (2024). AWS Lambda Developer Guide. https://docs.aws.amazon.com/lambda/ | '
    'Sbarski, P., & Kroonenburg, S. (2017). Serverless Architectures on AWS. Manning Publications.'
)
spacer()

# ── API Gateway ───────────────────────────────────────────────────────────────

h4('Amazon API Gateway')

body(
    'Amazon API Gateway es el único punto de entrada HTTP del backend. '
    'Gestiona y expone las rutas del sistema, enrutando cada solicitud a la función Lambda '
    'correspondiente y centralizando autenticación, rate limiting y logging sin requerir '
    'servidores de API propios. Esta separación de responsabilidades —API Gateway maneja '
    'la capa de transporte mientras Lambda maneja la lógica de negocio— sigue el principio '
    'de diseño de sistemas distribuidos que desacopla las preocupaciones de red '
    'de las de procesamiento (Richardson, 2018).'
)

ref_note(
    'Referencias: Amazon Web Services. (2024). Amazon API Gateway Developer Guide. '
    'https://docs.aws.amazon.com/apigateway/'
)
spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 2.4  METODOLOGÍA DE DESARROLLO
# ═══════════════════════════════════════════════════════════════════════════════

h2('2.4  Metodología de Desarrollo')

body(
    'En esta sección se describe el marco metodológico que guió el desarrollo '
    'del sistema a lo largo del proyecto.'
)

h3('2.4.1  SCRUM en equipos multidisciplinarios')

body(
    'Se explicará el framework SCRUM (Schwaber & Sutherland, 2020) y su aplicación '
    'en un equipo compuesto por perfiles de distintas disciplinas '
    '(ingeniería de software, diseño instruccional, expertos temáticos). '
    'Los puntos a cubrir son:'
)

scrum_points = [
    'Qué es SCRUM: framework ágil de desarrollo iterativo e incremental basado '
    'en sprints de duración fija, revisiones constantes y adaptación continua '
    '(Schwaber & Sutherland, 2020).',

    'Por qué SCRUM es adecuado para proyectos de IA: la calidad del output de un '
    'sistema LLM es desconocida hasta que se prueba con datos reales; los ciclos '
    'cortos permiten detectar y corregir problemas de prompt y arquitectura rápidamente.',

    'La dimensión multidisciplinaria: el equipo integró roles de desarrollo '
    '(backend, frontend) con roles instruccionales (LxD, Docente Revisor), '
    'lo que exige una coordinación explícita y roles bien definidos dentro del sprint.',

    'El rol del Docente Revisor como validador en cada sprint: '
    'cada iteración de prompts se valida con expertos reales, '
    'no solo con métricas automáticas.',
]

for pt in scrum_points:
    bullet(pt)

ref_note(
    'Referencias clave: Schwaber, K., & Sutherland, J. (2020). The Scrum Guide. Scrum.org. '
    'https://scrumguides.org/scrum-guide.html | '
    'Rigby, D. K., Sutherland, J., & Takeuchi, H. (2016). '
    'Embracing agile. Harvard Business Review, 94(5), 40-50.'
)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCIAS
# ═══════════════════════════════════════════════════════════════════════════════

h2('Referencias del Capítulo')

referencias = [
    'Anderson, L. W., & Krathwohl, D. R. (2001). A Taxonomy for Learning, Teaching, and Assessing: '
    "A Revision of Bloom's Taxonomy of Educational Objectives. Longman.",

    'Amazon Web Services. (2024). Amazon API Gateway Developer Guide. '
    'https://docs.aws.amazon.com/apigateway/',

    'Amazon Web Services. (2024). Amazon SQS Developer Guide. '
    'https://docs.aws.amazon.com/sqs/',

    'Amazon Web Services. (2024). Amazon SNS Developer Guide. '
    'https://docs.aws.amazon.com/sns/',

    'Amazon Web Services. (2024). AWS Lambda Developer Guide. '
    'https://docs.aws.amazon.com/lambda/',

    'Bloom, B. S., Engelhart, M. D., Furst, E. J., Hill, W. H., & Krathwohl, D. R. (1956). '
    'Taxonomy of Educational Objectives: The Classification of Educational Goals. David McKay.',

    'Brown, T. B., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., & Amodei, D. (2020). '
    'Language models are few-shot learners. Advances in Neural Information Processing Systems, 33, 1877-1901.',

    'Chase, H. (2024). LangGraph: Building stateful, multi-actor applications with LLMs. '
    'LangChain, Inc. https://langchain-ai.github.io/langgraph/',

    'Fielding, R. T., & Taylor, R. N. (2002). Principled design of the modern web architecture. '
    'ACM Transactions on Internet Technology, 2(2), 115-150.',

    'Kasneci, E., Sessler, K., Küchemann, S., Bannert, M., Dementieva, D., Fischer, F., & Kasneci, G. (2023). '
    'ChatGPT for good? On opportunities and challenges of large language models for education. '
    'Learning and Individual Differences, 103, 102274.',

    'Kleppmann, M. (2017). Designing Data-Intensive Applications: '
    'The Big Ideas Behind Reliable, Scalable, and Maintainable Systems. O\'Reilly Media.',

    'Liu, P., Yuan, W., Fu, J., Jiang, Z., Hayashi, H., & Neubig, G. (2023). '
    'Pre-train, prompt, and predict: A systematic survey of prompting methods in natural language processing. '
    'ACM Computing Surveys, 55(9), 1-35.',

    'OpenAI. (2024). Structured outputs. OpenAI Platform Documentation. '
    'https://platform.openai.com/docs/guides/structured-outputs',

    'Park, J. S., O\'Brien, J. C., Cai, C. J., Morris, M. R., Liang, P., & Bernstein, M. S. (2023). '
    'Generative agents: Interactive simulacra of human behavior. '
    'Proceedings of the 36th Annual ACM Symposium on User Interface Software and Technology (UIST).',

    'Raschka, S., Patterson, J., & Nolet, C. (2020). '
    'Machine learning in Python: Main developments and technology trends in data science, '
    'machine learning, and artificial intelligence. Information, 11(4), 193.',

    'Richardson, C. (2018). Microservices Patterns: With Examples in Java. Manning Publications.',

    'Rigby, D. K., Sutherland, J., & Takeuchi, H. (2016). '
    'Embracing agile. Harvard Business Review, 94(5), 40-50.',

    'Sbarski, P., & Kroonenburg, S. (2017). Serverless Architectures on AWS. Manning Publications.',

    'Schwaber, K., & Sutherland, J. (2020). The Scrum Guide. Scrum.org. '
    'https://scrumguides.org/scrum-guide.html',

    'Wei, J., Wang, X., Schuurmans, D., Bosma, M., Xia, F., Chi, E., & Zhou, D. (2022). '
    'Chain-of-thought prompting elicits reasoning in large language models. '
    'Advances in Neural Information Processing Systems, 35, 24824-24837.',

    'Wiggins, G., & McTighe, J. (2005). Understanding by Design (2nd ed.). ASCD.',

    'Wittig, M., & Wittig, A. (2018). Amazon Web Services in Action (2nd ed.). Manning Publications.',

    'Wu, Q., Bansal, G., Zhang, J., Wu, Y., Li, B., Zhu, E., & Wang, C. (2023). '
    'AutoGen: Enabling next-gen LLM applications via multi-agent conversation. '
    'arXiv preprint arXiv:2308.08155.',

    'Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K., & Cao, Y. (2023). '
    'ReAct: Synergizing reasoning and acting in language models. '
    'International Conference on Learning Representations (ICLR).',
]

for ref in referencias:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(ref)
    p.paragraph_format.space_after = Pt(3)

# ── Save ──────────────────────────────────────────────────────────────────────

out = 'D:/Proyectos actuales/proyecto_tesis_upc/Capitulo2_Marcoteorico_v3.docx'
doc.save(out)
print(f'Saved: {out}')
