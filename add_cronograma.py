"""
Add historical development cronograma to both documents.
Inserts a new table covering Sep 2025 – Apr 2026 (real work done)
before the existing planned-sprint table.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import copy

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ── Low-level cell builders (matching existing dark-blue style) ───────────────
def _qn(t): return f'{{{NS}}}{t}'

def hdr_cell(text, width):
    tc = etree.Element(_qn('tc'))
    tcPr = etree.SubElement(tc, _qn('tcPr'))
    tcW  = etree.SubElement(tcPr, _qn('tcW'));  tcW.set(_qn('w'), str(width)); tcW.set(_qn('type'), 'dxa')
    shd  = etree.SubElement(tcPr, _qn('shd'));  shd.set(_qn('val'),'clear'); shd.set(_qn('color'),'auto'); shd.set(_qn('fill'),'1F3864')
    p    = etree.SubElement(tc, _qn('p'))
    r    = etree.SubElement(p,  _qn('r'))
    rPr  = etree.SubElement(r,  _qn('rPr'))
    etree.SubElement(rPr, _qn('b'))
    col  = etree.SubElement(rPr, _qn('color')); col.set(_qn('val'), 'FFFFFF')
    t    = etree.SubElement(r,  _qn('t')); t.text = text
    return tc

def dat_cell(text, width, bold=False):
    tc = etree.Element(_qn('tc'))
    tcPr = etree.SubElement(tc, _qn('tcPr'))
    tcW  = etree.SubElement(tcPr, _qn('tcW')); tcW.set(_qn('w'), str(width)); tcW.set(_qn('type'), 'dxa')
    p    = etree.SubElement(tc, _qn('p'))
    # Handle multi-line text
    for i, line in enumerate(text.split('\n')):
        if i == 0:
            r = etree.SubElement(p, _qn('r'))
            if bold:
                rPr = etree.SubElement(r, _qn('rPr')); etree.SubElement(rPr, _qn('b'))
            t = etree.SubElement(r, _qn('t')); t.text = line
        else:
            br_p = etree.SubElement(tc, _qn('p'))
            r    = etree.SubElement(br_p, _qn('r'))
            if bold:
                rPr = etree.SubElement(r, _qn('rPr')); etree.SubElement(rPr, _qn('b'))
            t    = etree.SubElement(r, _qn('t')); t.text = line
    return tc

def make_tr(*cells):
    tr = etree.Element(_qn('tr'))
    for c in cells: tr.append(c)
    return tr

# ── Build the historical cronograma table ─────────────────────────────────────
def build_hist_table():
    # 5 cols: Sprint | Período | Fase | Actividades | Entregable
    W = [1600, 1500, 1200, 3200, 1500]

    tbl = etree.Element(_qn('tbl'))
    tblPr = etree.SubElement(tbl, _qn('tblPr'))
    ts = etree.SubElement(tblPr, _qn('tblStyle')); ts.set(_qn('val'), 'Tablaconcuadrcula')
    tw = etree.SubElement(tblPr, _qn('tblW'));     tw.set(_qn('w'), '0'); tw.set(_qn('type'), 'auto')
    tl = etree.SubElement(tblPr, _qn('tblLook'))
    for k, v in [('val','04A0'),('firstRow','1'),('lastRow','0'),('firstColumn','1'),
                 ('lastColumn','0'),('noHBand','0'),('noVBand','1')]:
        tl.set(_qn(k), v)
    tg = etree.SubElement(tbl, _qn('tblGrid'))
    for w in W:
        gc = etree.SubElement(tg, _qn('gridCol')); gc.set(_qn('w'), str(w))

    # Header
    headers = ['Sprint', 'Per\u00edodo', 'Fase', 'Actividades principales', 'Entregable / Resultado']
    tbl.append(make_tr(*[hdr_cell(h, W[i]) for i, h in enumerate(headers)]))

    rows = [
        (
            'Sprint 0\nOnboarding y dise\u00f1o',
            'Sep 16 \u2013 Sep 30, 2025',
            'Definici\u00f3n\ny dise\u00f1o',
            '- Recepci\u00f3n del repositorio del Tech Leader y onboarding: '
            'problema, arquitectura, stakeholders y alcance del sistema.\n'
            '- Identificaci\u00f3n de la brecha: contenido generado y validado, '
            'pero sin plataforma que gu\u00ede al usuario a lo largo del proceso IPES.\n'
            '- Dise\u00f1o inicial del frontend con flujo paso a paso '
            '(s\u00edlabo + kickoff como inputs).\n'
            '- Revisi\u00f3n de la arquitectura existente: LangGraph, FastAPI, agentes LLM.\n'
            '- Feedback del Tech Leader: priorizar implementaci\u00f3n de agentes antes del frontend.',
            'Comprensi\u00f3n del sistema heredado; plan de implementaci\u00f3n por casos de uso'
        ),
        (
            'Sprint 1\nImplementaci\u00f3n CU1\u2013CU3',
            'Oct 1 \u2013 Oct 31, 2025',
            'Implementaci\u00f3n',
            '- CU1: Carga de archivos de entrada (s\u00edlabo .xlsx + kickoff .docx).\n'
            '- CU2: Parsing de documentos \u2192 estructura JSON.\n'
            '- Tech Leader implementa CourseSchemaAgent \u2192 genera EsquemaCurso (1.\u00ba output).\n'
            '- CU3: Exportaci\u00f3n de EsquemaCurso + EsquemaActividades a Excel.\n'
            '- Ciclos de correcci\u00f3n, depuraci\u00f3n y ajustes de integraci\u00f3n.',
            'CU1, CU2 y CU3 funcionales;\nEsquemaCurso generado como primer output del sistema'
        ),
        (
            'Sprint 2\nImplementaci\u00f3n CU4 (IPES)',
            'Nov 1 \u2013 Nov 30, 2025',
            'Implementaci\u00f3n',
            '- CU4: Implementaci\u00f3n del IpesAgent '
            '(Introducci\u00f3n, Presentaciones y Ejercicios por semana).\n'
            '- Integraci\u00f3n del pipeline completo: Dise\u00f1o \u2192 Generaci\u00f3n IPES.\n'
            '- Pruebas con s\u00edlabos y kickoffs reales.\n'
            '- Ciclos intensivos de depuraci\u00f3n y correcci\u00f3n del flujo extremo a extremo.',
            'Pipeline completo funcional;\nCU4 (IpesAgent) operativo para las 18 semanas'
        ),
        (
            'Sprint 3\nEstabilizaci\u00f3n y validaci\u00f3n',
            'Dic 1 \u2013 Dic 31, 2025',
            'Correcci\u00f3n,\ncierre y\nvalidaci\u00f3n',
            '- Ciclo intensivo de correcciones y debugging del sistema.\n'
            '- Pruebas del pipeline completo con m\u00faltiples cursos reales.\n'
            '- Estabilizaci\u00f3n del sistema Experto Tem\u00e1tico v1.\n'
            '- 15 dic: Demo con coordinadora DAV (Direcci\u00f3n Acad\u00e9mica Virtual): '
            'present\u00f3 inter\u00e9s en la generaci\u00f3n IPES y en el flujo paso a paso de la plataforma; '
            'solicit\u00f3 auth, roles y gesti\u00f3n de usuarios para escalamiento.\n'
            '- Vacaciones de fin de a\u00f1o (15 dic \u2013 3 ene).',
            'Sistema v1 estabilizado;\nvalidaci\u00f3n con stakeholder DAV realizada'
        ),
        (
            'Sprint 4\nPulido con LXD',
            'Ene 4 \u2013 Feb 28, 2026',
            'Optimizaci\u00f3n\ny validaci\u00f3n\nLXD',
            '- Retomada el 4 de enero; inicio de colaboraci\u00f3n con coordinadora LXD.\n'
            '- 3\u20134 reuniones de revisi\u00f3n: ajuste de reglas y criterios did\u00e1cticos de calidad.\n'
            '- Refinamiento de prompts de IpesAgent y ContentGenerationAgent seg\u00fan criterios LXD.\n'
            '- Pruebas con cursos de distintas \u00e1reas y ciclos de refinamiento.\n'
            '- Mid-feb: jefe directo indica preparar paquete del proyecto para '
            'presentaci\u00f3n a gerencia; proyecto queda en espera.',
            'IPES alineado a est\u00e1ndares LXD;\npaquete preparado para escalamiento gerencial'
        ),
        (
            'Sprint 5\nCU5: Generaci\u00f3n de contenido',
            'Mar 1 \u2013 Mar 31, 2026',
            'Implementaci\u00f3n\ny validaci\u00f3n\ncontenido',
            '- Mid-mar: reuni\u00f3n con jefe: revisi\u00f3n del estado del proyecto; '
            'IPES validado, plataforma implementada, pero generaci\u00f3n de contenido pendiente.\n'
            '- Tarea asignada: implementar CU5 Generaci\u00f3n de Contenido; completado en 1 semana.\n'
            '- Segunda quincena: validaci\u00f3n con docentes revisores \u2014 '
            'contenido valorado como interesante con observaciones menores.\n'
            '- Continuaci\u00f3n del pulido de IPES y contenido.',
            'CU5 (ContentGenerationAgent) implementado y validado con docentes revisores'
        ),
        (
            'Sprint 6\nValidaci\u00f3n ejecutiva\ny CU6\u2013CU7',
            'Abr 1 \u2013 Abr 15, 2026',
            'Validaci\u00f3n\nexecutiva\n(en curso)',
            '- Pulido continuo de IPES y generaci\u00f3n de contenido.\n'
            '- Lun 13 abr: presentaci\u00f3n informal del proyecto por el PO ante directivos y CEO; '
            'propuesta bien recibida.\n'
            '- \u00daltimas validaciones en curso con usuarios clave.\n'
            '- Implementaci\u00f3n de CU6 (Agregar Referencias al contenido) y '
            'CU7 (Generaci\u00f3n de Consignas).\n'
            '- Elaboraci\u00f3n del Project Charter y Cap\u00edtulo 1 del trabajo de suficiencia.',
            'Presentaci\u00f3n ejecutiva realizada;\nCU6 y CU7 en implementaci\u00f3n'
        ),
    ]

    for row_data in rows:
        tbl.append(make_tr(*[dat_cell(row_data[i], W[i]) for i in range(5)]))

    return tbl


# ── Helper: make a paragraph element with optional style/bold ─────────────────
def make_para_el(text, style_val=None, bold=False, italic=False):
    p = OxmlElement('w:p')
    if style_val:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle'); pStyle.set(qn('w:val'), style_val)
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:   etree.SubElement(rPr, _qn('b'))
    if italic: etree.SubElement(rPr, _qn('i'))
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_before(body, ref_el, *elements):
    idx = list(body).index(ref_el)
    for i, el in enumerate(elements):
        body.insert(idx + i, el)


# ══════════════════════════════════════════════════════════════════════════════
# CAPITULO 1  – insert after Para 155 ("Cronograma") empty body
# ══════════════════════════════════════════════════════════════════════════════
doc = Document('D:/Proyectos actuales/proyecto_tesis_upc/Capitulo1_Tesis_Final.docx')
body = doc.element.body

# Find the empty paragraph right after the "Cronograma" heading (para 156)
# and the "Costos del proyecto" heading that follows (para 157)
# We insert: intro para + label para + hist table
# between para 156 and para 157

# Walk body elements to find the right anchor
paras_in_body = [el for el in body if el.tag.endswith('}p')]
cron_heading = None
costos_heading = None
for p_el in body:
    if p_el.tag.endswith('}p'):
        text = ''.join(t.text or '' for t in p_el.iter(_qn('t')))
        style = ''
        pPr = p_el.find(_qn('pPr'))
        if pPr is not None:
            ps = pPr.find(_qn('pStyle'))
            if ps is not None: style = ps.get(_qn('val'), '')
        is_heading = any(s in style for s in ('Heading', 'Ttulo', 'Titulo', 'tulo'))
        if 'Cronograma' in text and is_heading:
            cron_heading = p_el
        if 'Costos' in text and is_heading and cron_heading is not None:
            costos_heading = p_el
            break

print(f'Cronograma heading found: {cron_heading is not None}')
print(f'Costos heading found: {costos_heading is not None}')

# Elements to insert before "Costos" heading
intro_text = (
    'A continuaci\u00f3n se presenta el cronograma real de desarrollo del sistema '
    'Experto Tem\u00e1tico, desde el inicio del proyecto el 16 de septiembre de 2025 '
    'hasta la fecha actual. Las actividades est\u00e1n organizadas en sprints mensuales '
    'que reflejan el trabajo efectivamente realizado.'
)

caption_text = 'Tabla 5 Cronograma de desarrollo (Sep 2025 \u2013 Abr 2026)'

to_insert = [
    make_para_el(intro_text),
    make_para_el(caption_text, style_val='Caption'),
    build_hist_table(),
    make_para_el(''),   # spacer
]

insert_before(body, costos_heading, *to_insert)
doc.save('D:/Proyectos actuales/proyecto_tesis_upc/Capitulo1_Tesis_Final.docx')
print('Capitulo1 saved')


# ══════════════════════════════════════════════════════════════════════════════
# PROJECT CHARTER  – insert before existing sprint table (table 16)
# ══════════════════════════════════════════════════════════════════════════════
doc2 = Document('D:/Proyectos actuales/proyecto_tesis_upc/Project_Charter_Tesis.docx')
body2 = doc2.element.body

# Find the sprint table (table 16) and insert before it
sprint_tbl_el = doc2.tables[16]._tbl

subhead_hist = make_para_el(
    'Fase de Desarrollo (Sep 2025 \u2013 Abr 2026) \u2013 Historial real',
    bold=True
)
caption2 = make_para_el(
    'Tabla 12a Cronograma de desarrollo (Sep 2025 \u2013 Abr 2026)',
    style_val='Caption'
)
spacer1 = make_para_el('')
subhead_plan = make_para_el(
    'Cronograma de titulaci\u00f3n (May 2026 \u2013 Oct 2026) \u2013 Planificado',
    bold=True
)

to_insert2 = [subhead_hist, caption2, build_hist_table(), spacer1, subhead_plan]

idx2 = list(body2).index(sprint_tbl_el)
for i, el in enumerate(to_insert2):
    body2.insert(idx2 + i, el)

doc2.save('D:/Proyectos actuales/proyecto_tesis_upc/Project_Charter_Tesis.docx')
print('Project Charter saved')
